"""
Document Loader Module

Handles file type detection and loading for various document formats:
- PDF (with native text extraction and image rendering for OCR)
- DOCX (structured text extraction)
- TXT/Markdown (raw text)
- CSV/XLSX (table extraction as Markdown)
- Images (PNG, JPG, WEBP - passed to OCR)
"""

import hashlib
import logging
import mimetypes
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

logger = logging.getLogger(__name__)


class FileType(Enum):
    """Supported file types for ingestion."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    XLSX = "xlsx"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class PageContent:
    """Content from a single page or section of a document."""
    page_number: int
    text: str
    has_images: bool = False
    image_data: Optional[bytes] = None  # Raw image bytes for OCR if needed
    tables: List[str] = field(default_factory=list)  # Markdown tables
    is_scanned: bool = False  # True if page appears to be scanned/image-only


@dataclass
class LoadedDocument:
    """Result of loading a document."""
    file_path: str
    file_type: FileType
    file_hash: str  # SHA-256 of file content
    pages: List[PageContent]
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


def detect_file_type(file_path: Union[str, Path]) -> FileType:
    """
    Detect file type from extension and MIME type.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Detected FileType
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    # Extension-based detection
    ext_map = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".doc": FileType.DOCX,  # May need conversion
        ".txt": FileType.TXT,
        ".md": FileType.TXT,
        ".csv": FileType.CSV,
        ".xlsx": FileType.XLSX,
        ".xls": FileType.XLSX,
        ".png": FileType.IMAGE,
        ".jpg": FileType.IMAGE,
        ".jpeg": FileType.IMAGE,
        ".webp": FileType.IMAGE,
        ".tiff": FileType.IMAGE,
        ".tif": FileType.IMAGE,
        ".bmp": FileType.IMAGE,
    }
    
    if ext in ext_map:
        return ext_map[ext]
    
    # Fallback to MIME type
    mime_type, _ = mimetypes.guess_type(str(path))
    if mime_type:
        if "pdf" in mime_type:
            return FileType.PDF
        if "word" in mime_type or "document" in mime_type:
            return FileType.DOCX
        if "text" in mime_type:
            return FileType.TXT
        if "csv" in mime_type:
            return FileType.CSV
        if "spreadsheet" in mime_type or "excel" in mime_type:
            return FileType.XLSX
        if "image" in mime_type:
            return FileType.IMAGE
    
    return FileType.UNKNOWN


def _compute_file_hash(file_path: Path) -> str:
    """Compute SHA-256 hash of file content."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256_hash.update(chunk)
    return sha256_hash.hexdigest()


class DocumentLoader:
    """
    Load documents from various file formats.
    
    Supports PDF, DOCX, TXT, CSV/XLSX, and images.
    """
    
    def __init__(
        self,
        pdf_dpi: int = 300,
        extract_tables: bool = True,
        max_pages: Optional[int] = None,
    ):
        """
        Initialize the document loader.
        
        Args:
            pdf_dpi: DPI for PDF image rendering (default 300)
            extract_tables: Whether to extract tables as Markdown
            max_pages: Maximum pages to process (None for all)
        """
        self.pdf_dpi = pdf_dpi
        self.extract_tables = extract_tables
        self.max_pages = max_pages
        
    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        """
        Load a document from file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            LoadedDocument with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            return LoadedDocument(
                file_path=str(path),
                file_type=FileType.UNKNOWN,
                file_hash="",
                pages=[],
                error=f"File not found: {path}",
            )
        
        file_type = detect_file_type(path)
        file_hash = _compute_file_hash(path)
        
        try:
            if file_type == FileType.PDF:
                return self._load_pdf(path, file_hash)
            elif file_type == FileType.DOCX:
                return self._load_docx(path, file_hash)
            elif file_type == FileType.TXT:
                return self._load_txt(path, file_hash)
            elif file_type == FileType.CSV:
                return self._load_csv(path, file_hash)
            elif file_type == FileType.XLSX:
                return self._load_xlsx(path, file_hash)
            elif file_type == FileType.IMAGE:
                return self._load_image(path, file_hash)
            else:
                return LoadedDocument(
                    file_path=str(path),
                    file_type=file_type,
                    file_hash=file_hash,
                    pages=[],
                    error=f"Unsupported file type: {path.suffix}",
                )
        except Exception as e:
            logger.error(f"Error loading document {path}: {e}")
            return LoadedDocument(
                file_path=str(path),
                file_type=file_type,
                file_hash=file_hash,
                pages=[],
                error=str(e),
            )
    
    def _load_pdf(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load PDF document with text extraction and image rendering."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to basic PDF loading")
            return self._load_pdf_basic(path, file_hash)
        
        pages = []
        doc = fitz.open(str(path))
        
        try:
            num_pages = min(len(doc), self.max_pages) if self.max_pages else len(doc)
            
            for page_num in range(num_pages):
                page = doc[page_num]
                text = page.get_text()
                
                # Check if page is scanned (low text density relative to page size)
                text_density = len(text.strip()) / (page.rect.width * page.rect.height + 1)
                is_scanned = text_density < 0.001  # Very low text density
                
                # Extract tables if present
                tables = []
                if self.extract_tables:
                    try:
                        table_finder = page.find_tables()
                        for table in table_finder:
                            md_table = self._table_to_markdown(table.extract())
                            if md_table:
                                tables.append(md_table)
                    except Exception as e:
                        logger.debug(f"Table extraction failed for page {page_num}: {e}")
                
                # Render page as image if scanned or has low text
                image_data = None
                if is_scanned or len(text.strip()) < 100:
                    pixmap = page.get_pixmap(dpi=self.pdf_dpi)
                    image_data = pixmap.tobytes("png")
                
                pages.append(PageContent(
                    page_number=page_num + 1,
                    text=text,
                    has_images=image_data is not None,
                    image_data=image_data,
                    tables=tables,
                    is_scanned=is_scanned,
                ))
        finally:
            doc.close()
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.PDF,
            file_hash=file_hash,
            pages=pages,
            metadata={"page_count": len(pages)},
        )
    
    def _load_pdf_basic(self, path: Path, file_hash: str) -> LoadedDocument:
        """Basic PDF loading without PyMuPDF (fallback)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return LoadedDocument(
                file_path=str(path),
                file_type=FileType.PDF,
                file_hash=file_hash,
                pages=[],
                error="No PDF library available (install pymupdf or pypdf)",
            )
        
        reader = PdfReader(str(path))
        pages = []
        
        num_pages = min(len(reader.pages), self.max_pages) if self.max_pages else len(reader.pages)
        
        for i in range(num_pages):
            page = reader.pages[i]
            text = page.extract_text() or ""
            pages.append(PageContent(
                page_number=i + 1,
                text=text,
                is_scanned=len(text.strip()) < 50,
            ))
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.PDF,
            file_hash=file_hash,
            pages=pages,
            metadata={"page_count": len(pages)},
        )
    
    def _load_docx(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load DOCX document with structure preservation."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return LoadedDocument(
                file_path=str(path),
                file_type=FileType.DOCX,
                file_hash=file_hash,
                pages=[],
                error="python-docx not installed",
            )
        
        doc = DocxDocument(str(path))
        
        # Extract text with paragraph structure
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            if para.style and para.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}\n")
            else:
                paragraphs.append(text)
        
        # Extract tables
        tables = []
        if self.extract_tables:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    md_table = self._rows_to_markdown(rows)
                    tables.append(md_table)
        
        full_text = "\n".join(paragraphs)
        
        # DOCX is treated as a single page
        pages = [PageContent(
            page_number=1,
            text=full_text,
            tables=tables,
        )]
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.DOCX,
            file_hash=file_hash,
            pages=pages,
            metadata={"paragraph_count": len(paragraphs)},
        )
    
    def _load_txt(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load plain text file."""
        text = path.read_text(encoding="utf-8", errors="replace")
        
        pages = [PageContent(
            page_number=1,
            text=text,
        )]
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.TXT,
            file_hash=file_hash,
            pages=pages,
            metadata={"char_count": len(text)},
        )
    
    def _load_csv(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load CSV file as Markdown table."""
        import csv
        
        rows = []
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)
        
        md_table = self._rows_to_markdown(rows) if rows else ""
        
        pages = [PageContent(
            page_number=1,
            text=md_table,
            tables=[md_table] if md_table else [],
        )]
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.CSV,
            file_hash=file_hash,
            pages=pages,
            metadata={"row_count": len(rows)},
        )
    
    def _load_xlsx(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load Excel file as Markdown tables (one per sheet)."""
        try:
            import openpyxl
        except ImportError:
            return LoadedDocument(
                file_path=str(path),
                file_type=FileType.XLSX,
                file_hash=file_hash,
                pages=[],
                error="openpyxl not installed",
            )
        
        wb = openpyxl.load_workbook(str(path), data_only=True)
        pages = []
        
        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            rows = []
            
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cells):  # Skip empty rows
                    rows.append(cells)
            
            if rows:
                md_table = self._rows_to_markdown(rows)
                text = f"## Sheet: {sheet_name}\n\n{md_table}"
                
                pages.append(PageContent(
                    page_number=sheet_idx + 1,
                    text=text,
                    tables=[md_table],
                ))
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.XLSX,
            file_hash=file_hash,
            pages=pages,
            metadata={"sheet_count": len(wb.sheetnames)},
        )
    
    def _load_image(self, path: Path, file_hash: str) -> LoadedDocument:
        """Load image file (will need OCR processing)."""
        image_data = path.read_bytes()
        
        pages = [PageContent(
            page_number=1,
            text="",  # To be filled by OCR
            has_images=True,
            image_data=image_data,
            is_scanned=True,
        )]
        
        return LoadedDocument(
            file_path=str(path),
            file_type=FileType.IMAGE,
            file_hash=file_hash,
            pages=pages,
            metadata={"image_size": len(image_data)},
        )
    
    def _table_to_markdown(self, table_data: List[List[Any]]) -> str:
        """Convert extracted table data to Markdown format."""
        if not table_data:
            return ""
        
        # Clean cells
        rows = []
        for row in table_data:
            cells = [str(cell).strip() if cell else "" for cell in row]
            rows.append(cells)
        
        return self._rows_to_markdown(rows)
    
    def _rows_to_markdown(self, rows: List[List[str]]) -> str:
        """Convert list of rows to Markdown table."""
        if not rows:
            return ""
        
        # Determine column widths
        num_cols = max(len(row) for row in rows)
        
        # Normalize row lengths
        normalized = []
        for row in rows:
            row_copy = list(row)
            while len(row_copy) < num_cols:
                row_copy.append("")
            normalized.append(row_copy)
        
        # Build Markdown table
        lines = []
        
        # Header
        if normalized:
            header = "| " + " | ".join(normalized[0]) + " |"
            separator = "| " + " | ".join(["---"] * num_cols) + " |"
            lines.append(header)
            lines.append(separator)
        
        # Data rows
        for row in normalized[1:]:
            line = "| " + " | ".join(row) + " |"
            lines.append(line)
        
        return "\n".join(lines)
